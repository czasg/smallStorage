from docx import Document
import numpy as np


def get_document(file):
    return Document(str)


class _Table(object):
    """
    主要提供merged_cells属性
    思路：
    1、按行转化为矩阵，合并单元格内存地址相同的，依次标记为1,2,3...
    2、按列转化为矩阵，合并单元格内存地址相同的，依次标记为1,2,3...
    3、将上述两矩阵相乘，获取融合矩阵，判断是否存在跨越行列的单元格，存在则按下标取出，
       取出样式为 [(左上角行下标,右下角行下标,左上角列下标,右上角列下标)]
    4、获取行矩阵中的合并单元格
       取出样式为 [(左上角行下标,右下角行下标,左上角列下标,右上角列下标)]
    5、获取列矩阵中的合并单元格
       取出样式为 [(左上角行下标,右下角行下标,左上角列下标,右上角列下标)]
    6、检查行、列矩阵中是否存在与融合矩阵中重合的单元格，去重
    """
    def __init__(self, table):
        self.table = table
        self.nrows = len(self.table.column_cells(0))  # 总行数
        self.ncols = len(self.table.row_cells(0))  # 总列数

    @property
    def merged_cells(self):
        # 获取行矩阵
        row_mat = self.get_matrix(row=True)
        # 获取列矩阵
        col_mat = self.get_matrix(col=True).T
        # 获取融合后的合并单元格
        fusion_res = self.get_fusion_merge_cells(row_mat, col_mat)  # todo
        # 获取行矩阵中的合并单元格
        row_res = self.get_row_merge_cells(row_mat)  # def func  # todo
        # 获取列矩阵中的合并单元格
        col_res = self.get_col_merge_cells(col_mat)  # def func  # todo
        # 检查行列单元格中的重复，并返回结果
        return self.chech_result_merge_cells(fusion_res, row_res, col_res)  # list

    def get_matrix(self, row=False, col=False):
        father_num, child_num = 0, 0
        get_data = lambda x: x

        if row:
            father_num = self.nrows
            child_num = self.ncols
            get_data = self.table.row_cells
        elif col:
            father_num = self.ncols
            child_num = self.nrows
            get_data = self.table.column_cells
        res_list = []

        for i in range(father_num):
            all_cells = get_data(i)
            all_set = set(all_cells)
            res = np.zeros(child_num)
            if len(all_set) == child_num:
                res_list.append(res)
                continue
            pool = {}
            index_pool = {}
            flag = 1
            for i, cell in enumerate(all_cells):
                if cell in pool:
                    if pool[cell] == 1:
                        res[i] += flag
                        res[i - 1] += flag
                        index_pool[cell] = flag  #
                        pool[cell] += 1  #
                        flag += 1
                    elif pool[cell] > 1:
                        res[i] = index_pool[cell]
                    continue
                pool.setdefault(cell, 1)
            res_list.append(res)
        return np.array(res_list)

    def get_fusion_merge_cells(self, row_mat, col_mat):
        fusion = row_mat * col_mat
        print(row_mat)
        print(col_mat)
        print(fusion)
        res = []
        index = 1
        while True:
            if index in fusion:
                ix = np.isin(fusion, index)
                x, y = np.where(ix)
                cursor_list = list(zip(x, y))  #
                m, n = cursor_list[0]
                start, end = m, n
                while True:
                    n += 1  # right sep
                    if (m, n) in cursor_list:
                        m += 1
                        if (m, n) in cursor_list:
                            continue
                        m -= 1  # restore
                        continue
                    n -= 1  # restore
                    m += 1
                    if (m, n) in cursor_list:
                        continue
                    m -= 1
                    res.append((start, m, end, n))
                    next_cursor = cursor_list[:]
                    for data in next_cursor:
                        if data == (m, n):
                            cursor_list.remove(data)
                            if cursor_list:
                                m, n = cursor_list[0]
                                start, end = m, n
                                break
                            else:
                                break
                        cursor_list.remove(data)
                    if not cursor_list:
                        break
                index += 1
            else:
                break
        return res

    def get_row_merge_cells(self, row_mat):
        res = []
        index = 1
        while True:
            if index in row_mat:
                ix = np.isin(row_mat, index)
                x, y = np.where(ix)
                cursor_list = list(zip(x, y))  #
                m, n = cursor_list[0]
                start, end = m, n
                while True:
                    n += 1  # right sep
                    if (m, n) in cursor_list:
                        continue
                    n -= 1
                    res.append((start, m, end, n))
                    next_cursor = cursor_list[:]
                    for data in next_cursor:
                        if data == (m, n):
                            cursor_list.remove(data)
                            if cursor_list:
                                m, n = cursor_list[0]
                                start, end = m, n
                                break
                            else:
                                break
                        cursor_list.remove(data)
                    if not cursor_list:
                        break
                index += 1
            else:
                break
        return res

    def get_col_merge_cells(self, col_mat):
        res = []
        index = 1
        while True:
            if index in col_mat:
                ix = np.isin(col_mat, index)
                x, y = np.where(ix)
                cursor_list = list(zip(x, y))  #
                m, n = cursor_list[0]
                start, end = m, n
                while True:
                    m += 1  # right sep
                    if (m, n) in cursor_list:
                        continue
                    m -= 1
                    res.append((start, m, end, n))
                    next_cursor = cursor_list[:]
                    for data in next_cursor:
                        if data == (m, n):
                            cursor_list.remove(data)
                            if cursor_list:
                                m, n = cursor_list[0]
                                start, end = m, n
                                break
                            else:
                                break
                        cursor_list.remove(data)
                    if not cursor_list:
                        break
                index += 1
            else:
                break
        return res

    def chech_result_merge_cells(self, fusion_res, row_res, col_res):
        res = fusion_res[:]
        r_res = row_res[:]
        c_res = col_res[:]
        if fusion_res:
            for x1, y1, m1, n1 in row_res:
                for x, y, m, n in fusion_res:
                    if x <= x1 <= y and x <= y1 <= y and m <= m1 <= n and m <= n1 <= n:
                        r_res.remove((x1, y1, m1, n1))
                        continue
            for x1, y1, m1, n1 in col_res:
                for x, y, m, n in fusion_res:
                    if x <= x1 <= y and x <= y1 <= y and m <= m1 <= n and m <= n1 <= n:
                        c_res.remove((x1, y1, m1, n1))
                        continue
        return res + r_res + c_res


class DocxReader(object):
    def __init__(self, document):
        self.document = document
        self.tables = self.document.tables
        self.tables_res = []

    @classmethod
    def from_str(cls, string=None, path=None):
        file = string if string else path
        if file:
            return cls(get_document(file))
        else:
            raise Exception("None thing to reading")

    def table_manage(self, ):
        pass

    def table2dict(self, ):
        pass


if __name__ == "__main__":
    # doc = Document('text.docx')
    doc = Document('test1.docx')
    # doc = Document('test2.docx')
    table = doc.tables[0]
    print(_Table(table).merged_cells)
