from docx import Document
"""
主要尝试一些写的操作，pass
"""
def test():
    doc = Document()  # 'text.docx'
    # doc.add_heading("czaOrz", 0)
    # doc.add_paragraph('hello?cza')
    table = doc.add_table(rows=3, cols=3)
    rows = table.rows
    rows[0].cells[0].text = '1'
    rows[0].cells[1].text = '2'
    rows[0].cells[2].text = '3'

    rows[1].cells[0].add_table(rows=2, cols=2)


    doc.save('text.docx')  # save in a new docx file


if __name__ == "__main__":
    test()