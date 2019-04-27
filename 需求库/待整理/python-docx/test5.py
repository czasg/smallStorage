from docx import Document
import numpy as np  # np.array(list)

doc = Document('text.docx')
# doc = Document('test1.docx')
table = doc.tables[0]
nrows = len(table.column_cells(0))
ncols = len(table.row_cells(0))  # print(nrows,ncols) 4-3
# print(nrows,ncols)
res_list = []
for i in range(nrows):  #
    row_cells = table.row_cells(i)
    row_set = set(row_cells)
    res = np.zeros(ncols)
    if len(row_set) == ncols:
        res_list.append(res)
        continue
    pool = {}
    index_pool = {}
    flag = 1
    for i,cell in enumerate(row_cells):
        if cell in pool:
            if pool[cell] == 1:
                res[i] += flag
                res[i-1] += flag
                index_pool[cell] = flag  #
                pool[cell] += 1  #
                flag += 1
            elif pool[cell] > 1:
                res[i] = index_pool[cell]
            continue
        pool.setdefault(cell, 1)  # append(cell)
    res_list.append(res)
row_matrix = np.array(res_list)


res_list = []
for i in range(ncols):  #
    col_cells = table.column_cells(i)
    col_set = set(col_cells)
    res = np.zeros(nrows)
    if len(col_set) == nrows:
        res_list.append(res)
        continue
    pool = []
    flag = 1
    for i,cell in enumerate(col_cells):
        if cell in pool:
            res[i] += flag
            res[i-1] += flag
            flag += 1
            continue
        pool.append(cell)
    res_list.append(res)
col_matrix = np.array(res_list).T
# print(col_matrix)

res = row_matrix*col_matrix

# todo check the long condition
index = 1
cza = []
while True:
    if index in res:
        ix = np.isin(res, index)
        x,y = np.where(ix)
        # print(x,y)
        aaaaa = list(zip(x,y))
        # print("aaaaa", aaaaa)
        # bbbbb = np.array(aaaaa)
        # print(bbbbb)
        m,n = aaaaa[0]
        start = m
        end = n
        while True:
            n += 1  # right sep
            if (m,n) in aaaaa:
                m += 1
                if (m,n) in aaaaa:
                    continue
                m -= 1  # restore
                continue
            n -= 1  # restore
            m += 1
            if (m, n) in aaaaa:
                continue
            m -= 1
            cza.append((start, m, end, n))
            ttttttt = aaaaa[:]
            for data in ttttttt:
                if data == (m,n):
                    aaaaa.remove(data)
                    if aaaaa:
                        m, n = aaaaa[0]
                        start = m
                        end = n
                        break
                    else:
                        break
                aaaaa.remove(data)
            if not aaaaa:
                break
        index += 1
    else:
        break
# print(cza)

print(res, )
print(np.where(res>0))
x,y = np.where(res>0)
cursor_list = list(zip(x, y))
m, n = cursor_list[0]
start, end = m, n
while True:
    n += 1  # right sep
    if (m, n) in cursor_list:
        m += 1
        if (m, n) in cursor_list:
            continue
        m -= 1  # restore
        continue
    n -= 1  # restore
    m += 1
    if (m, n) in cursor_list:
        continue
    m -= 1
    cza.append((start, m, end, n))
    next_cursor = cursor_list[:]
    for data in next_cursor:
        if data == (m, n):
            cursor_list.remove(data)
            if cursor_list:
                m, n = cursor_list[0]
                start, end = m, n
                break
            else:
                break
        cursor_list.remove(data)
    if not cursor_list:
        break
print(cza)
