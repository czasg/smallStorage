from docx import Document

def test():
    doc = Document('database.docx')
    index,key,value = [],[],[]
    table_index = 0
    for table in doc.tables:
        table_index += 1
        row_index = 0
        for row in table.rows:
            row_index += 1
            list = []
            for cell in row.cells:
                text = ''
                # for p in cell.paragraphs:
                #     text += p.text
                text += cell.text
                if row_index == 1:
                    key.append(text)
                else:
                    list.append(text)
            a = value.append(list) if list else list
    print(key)
    print(value)

if __name__ == "__main__":
    test()